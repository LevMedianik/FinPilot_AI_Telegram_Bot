import os
import fitz
import docx
from uuid import uuid4
from dotenv import load_dotenv
import requests
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.docstore.document import Document

load_dotenv()

# ----------------------------
# Storage (files + FAISS index)
# ----------------------------
DATA_DIR = "./data"
INDEX_DIR = "./faiss_index"
os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(INDEX_DIR, exist_ok=True)

def save_file(file_bytes, original_filename: str) -> str:
    file_id = str(uuid4())
    filename = f"{file_id}_{original_filename}"
    filepath = os.path.join(DATA_DIR, filename)
    with open(filepath, "wb") as f:
        f.write(file_bytes)
    return filepath

# ----------------------------
# Text extraction
# ----------------------------
def extract_text_from_pdf(filepath: str) -> str:
    doc = fitz.open(filepath)
    return "".join(page.get_text() for page in doc)

def extract_text_from_docx(filepath: str) -> str:
    doc = docx.Document(filepath)
    return "\n".join(para.text for para in doc.paragraphs)

def extract_text_from_txt(filepath: str) -> str:
    with open(filepath, "r", encoding="utf-8") as f:
        return f.read()

def extract_text_from_file(filepath: str) -> str:
    lower = filepath.lower()
    if lower.endswith(".pdf"):
        return extract_text_from_pdf(filepath)
    if lower.endswith(".docx"):
        return extract_text_from_docx(filepath)
    if lower.endswith(".txt"):
        return extract_text_from_txt(filepath)
    raise ValueError("❌ Поддерживаются только PDF, DOCX и TXT файлы.")

# ----------------------------
# Vector index (FAISS)
# ----------------------------
_EMBED_MODEL = os.getenv("EMBED_MODEL", "sentence-transformers/all-MiniLM-L6-v2")

def index_text_with_faiss(text: str):
    splitter = RecursiveCharacterTextSplitter(chunk_size=300, chunk_overlap=50)
    chunks = splitter.split_text(text)
    documents = [Document(page_content=chunk) for chunk in chunks]

    embeddings = HuggingFaceEmbeddings(model_name=_EMBED_MODEL)
    vectorstore = FAISS.from_documents(documents, embedding=embeddings)
    vectorstore.save_local(INDEX_DIR)
    return vectorstore

def load_existing_index():
    index_file = os.path.join(INDEX_DIR, "index.faiss")
    if not os.path.exists(index_file):
        return None
    embeddings = HuggingFaceEmbeddings(model_name=_EMBED_MODEL)
    return FAISS.load_local(INDEX_DIR, embeddings, allow_dangerous_deserialization=True)

# ----------------------------
# OpenRouter LLM (robust routing)
# ----------------------------
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")

DEFAULT_MODELS = [
    "deepseek/deepseek-chat-v3-0324",  # stable id (may require credits)
    "deepseek/deepseek-chat",          # alias / fallback
]

def _models_from_env():
    raw = os.getenv("OPENROUTER_MODELS", "").strip()
    if not raw:
        return DEFAULT_MODELS
    return [m.strip() for m in raw.split(",") if m.strip()]

OPENROUTER_MODELS = _models_from_env()

HEADERS = {
    "Authorization": f"Bearer {OPENROUTER_API_KEY}" if OPENROUTER_API_KEY else "",
    "Content-Type": "application/json",
    "HTTP-Referer": os.getenv("OPENROUTER_REFERER", "http://localhost"),
    "X-Title": os.getenv("OPENROUTER_APP_NAME", "BizSense"),
}

OPENROUTER_URL = os.getenv("OPENROUTER_URL", "https://openrouter.ai/api/v1/chat/completions")

def call_openrouter(messages, temperature: float = 0.3, max_tokens: int | None = None) -> str:
    """
    Tries models in OPENROUTER_MODELS until one responds successfully.
    Returns assistant text or a readable error message.
    """
    if not OPENROUTER_API_KEY:
        return "Ошибка: не задан OPENROUTER_API_KEY в переменных окружения."

    last_err = None
    for model in OPENROUTER_MODELS:
        body = {"model": model, "messages": messages, "temperature": temperature}
        if max_tokens is not None:
            body["max_tokens"] = max_tokens

        try:
            resp = requests.post(OPENROUTER_URL, headers=HEADERS, json=body, timeout=60)
        except Exception as e:
            last_err = f"Ошибка запроса к LLM: {e}"
            continue

        if resp.status_code == 200:
            data = resp.json()
            try:
                return data["choices"][0]["message"]["content"]
            except Exception:
                return f"Ошибка разбора ответа LLM: {data}"

        try:
            err_json = resp.json()
        except Exception:
            err_json = {"raw": resp.text}

        msg = err_json.get("error", {}).get("message") or str(err_json)
        last_err = f"Ошибка запроса к LLM ({resp.status_code}) для модели {model}: {msg}"

        if resp.status_code == 404 and "No endpoints found" in msg:
            continue  # try next model

        break

    return last_err or "Не удалось получить ответ от LLM."

DEFAULT_SYSTEM_PROMPT = os.getenv(
    "SYSTEM_PROMPT",
    "Ты – ассистент для ответов по документам. "
    "Отвечай ТОЛЬКО фактами из предоставленного контекста. "
    "НЕ добавляй выводы, обобщения и интерпретации. "
    "НЕ используй фразы 'на основе контекста', 'можно выделить', 'вероятно', 'скорее всего'. "
    "Если в контексте нет прямого ответа – скажи: 'В документах это не найдено'. "
    "Формат: 2–6 коротких пунктов или 2–4 предложения. Без лишнего."
)

# --- RAG gating (anti-hallucination) ---
RAG_MIN_CONTEXT_CHARS = int(os.getenv("RAG_MIN_CONTEXT_CHARS", "300"))
RAG_MAX_L2_DISTANCE = float(os.getenv("RAG_MAX_L2_DISTANCE", "1.1"))
RAG_REFUSAL_TEXT = os.getenv(
    "RAG_REFUSAL_TEXT",
    "В текущей базе документов нет информации для ответа на этот вопрос.\n"
    "Я отвечаю только на основе загруженных материалов. "
    "Уточните запрос или загрузите другой документ."
)

def _normalize_ru(s: str) -> str:
    s = s.lower()
    for ch in ",.;:!?()[]{}\"'«»––-":
        s = s.replace(ch, " ")
    return " ".join(s.split())

def _has_overlap(question: str, context: str, min_hits: int = 1) -> bool:
    q = _normalize_ru(question)
    c = _normalize_ru(context)

    # базовые стоп-слова (минимально)
    stop = {"что", "как", "почему", "зачем", "где", "когда", "для", "в", "на", "и", "или", "это", "такие"}
    q_words = [w for w in q.split() if len(w) >= 5 and w not in stop]

    hits = sum(1 for w in set(q_words) if w in c)
    return hits >= min_hits

def _build_strict_rag_prompt(context: str, question: str) -> str:
    return (
        "Контекст (выдержки из документов):\n"
        f"{context}\n"
        "---\n"
        f"Вопрос: {question}\n\n"
        "Ответь ТОЛЬКО тем, что явно есть в контексте.\n"
        "Формат ответа (строго):\n"
        "Пункт: <краткий ответ>. Цитата: \"<дословная фраза из контекста>\"\n"
        "Если не можешь дать цитату – ответь: \"В документах это не найдено.\"\n"
        "Ответ:"
    )

def query_index(question: str, announce: bool = False):
    vectorstore = load_existing_index()
    if not vectorstore:
        return "❌ База знаний не найдена. Пожалуйста, загрузите документ."

    # 1) Достаём релевантные фрагменты вместе со score (distance)
    # Для FAISS в LangChain обычно это L2-distance: чем меньше, тем лучше.
    try:
        hits = vectorstore.similarity_search_with_score(question, k=8)
    except Exception:
        # fallback, если у твоей версии нет with_score
        retriever = vectorstore.as_retriever(search_type="mmr", search_kwargs={"k": 4})
        results = retriever.get_relevant_documents(question)
        context = "\n".join(doc.page_content.strip() for doc in results).strip()
        if len(context) < RAG_MIN_CONTEXT_CHARS:
            return (("🔍 Ищу ответ...", RAG_REFUSAL_TEXT) if announce else RAG_REFUSAL_TEXT)
        
        # если контекст есть – идём в LLM
        full_prompt = _build_strict_rag_prompt(context, question)
        messages = [
            {"role": "system", "content": DEFAULT_SYSTEM_PROMPT},
            {"role": "user", "content": full_prompt},
        ]
        reply = call_openrouter(messages=messages, temperature=0.3)
        return ("🔍 Ищу ответ...", reply) if announce else reply


    # 2) Сбор контекста + проверка качества
    if not hits:
        return (("🔍 Ищу ответ...", RAG_REFUSAL_TEXT) if announce else RAG_REFUSAL_TEXT)

    docs = [doc for (doc, _score) in hits]
    scores = [_score for (_doc, _score) in hits]

    context = "\n".join(d.page_content.strip() for d in docs).strip()
    best_score = min(scores) if scores else 999.0

    # Gate A: контекст слишком короткий (обычно означает "не нашлось")
    # Gate B: даже лучший score слабый (далеко от вопроса)
    if len(context) < RAG_MIN_CONTEXT_CHARS or best_score > RAG_MAX_L2_DISTANCE:
        return (("🔍 Ищу ответ...", RAG_REFUSAL_TEXT) if announce else RAG_REFUSAL_TEXT)
    min_hits = 0 if len(question.strip()) < 35 else 1
    if not _has_overlap(question, context, min_hits=min_hits):
        return (("🔍 Ищу ответ...", RAG_REFUSAL_TEXT) if announce else RAG_REFUSAL_TEXT)

    # 3) Есть релевантный контекст – зовём LLM
    full_prompt = _build_strict_rag_prompt(context, question)
    messages = [
        {"role": "system", "content": DEFAULT_SYSTEM_PROMPT},
        {"role": "user", "content": full_prompt},
    ]
    
    reply = call_openrouter(messages=messages, temperature=0.3)
    return ("🔍 Ищу ответ...", reply) if announce else reply

def summarize_pdf(announce: bool = False):
    vectorstore = load_existing_index()
    if not vectorstore:
        return "❌ Индекс не найден. Пожалуйста, загрузите документ."

    retriever = vectorstore.as_retriever(search_type="mmr", search_kwargs={"k": 6})
    results = retriever.get_relevant_documents("Основное содержание документа, тезисы, выводы")

    text = "\n".join(doc.page_content.strip() for doc in results)
    prompt = (
        "Сделай краткое резюме документа в 7–12 пунктов. "
        "Укажи ключевые идеи, определения и выводы.\n\n"
        f"{text}"
    )

    messages = [
        {"role": "system", "content": DEFAULT_SYSTEM_PROMPT},
        {"role": "user", "content": prompt},
    ]
    reply = call_openrouter(messages=messages, temperature=0.3)
    return ("📖 Пересказываю текст...", reply) if announce else reply
